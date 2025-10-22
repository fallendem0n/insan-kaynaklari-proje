import customtkinter as ctk
from tkinter import filedialog, messagebox
import pandas as pd
import docx 
import os
import re 
import traceback 

class MailMergerFrame(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master)

        self.excel_path = None
        self.word_path = None
        self.excel_columns = []
        self.placeholders = [] 
        self.naming_columns = [] 

        self.create_widgets()

    def create_widgets(self):
        self.info_textbox = ctk.CTkTextbox(self, height=100, wrap="word")
        self.info_textbox.pack(padx=10, pady=(10, 5), fill="x", expand=False)

        bilgi_metni = """
Bu araç, bir Excel veri dosyasındaki bilgileri bir Word şablonuna yerleştirir.

1. Excel veri kaynağınızı (.xlsx veya .xls) seçin.
2. `.docx` formatındaki Word şablonunuzu seçin.
3. Word şablonunuzdaki değiştirilecek alanlar `{{Sutun_Adi}}` formatında olmalıdır. (Çift süslü parantez)
4. Excel dosyanızdaki sütun başlığı (`Sutun_Adi`) ile Word'deki yer tutucu adı birebir aynı olmalıdır (büyük/küçük harf duyarlı).
5. Araç, Excel'deki her satır için ayrı bir Word belgesi oluşturacaktır.
        """
        self.info_textbox.configure(state="normal")
        self.info_textbox.insert("1.0", bilgi_metni.strip())
        self.info_textbox.configure(state="disabled")

        main_frame = ctk.CTkFrame(self, corner_radius=10)
        main_frame.pack(pady=10, padx=20, fill="both", expand=True)

        title_label = ctk.CTkLabel(main_frame, text="Word Taslak Oluşturucu", font=ctk.CTkFont(size=20, weight="bold"))
        title_label.pack(pady=(10, 10))

        excel_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        excel_frame.pack(pady=5, padx=10, fill="x")

        excel_button = ctk.CTkButton(excel_frame, text="1. Excel Veri Dosyası Seç", width=180, command=self.select_excel)
        excel_button.pack(side="left", padx=(0, 10))

        self.excel_label = ctk.CTkLabel(excel_frame, text="Veri dosyası seçilmedi...", text_color="gray")
        self.excel_label.pack(side="left", fill="x", expand=True)

        word_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        word_frame.pack(pady=5, padx=10, fill="x")

        word_button = ctk.CTkButton(word_frame, text="2. Word Şablon Dosyası Seç", width=180, command=self.select_word)
        word_button.pack(side="left", padx=(0, 10))

        self.word_label = ctk.CTkLabel(word_frame, text="Şablon dosyası seçilmedi...", text_color="gray")
        self.word_label.pack(side="left", fill="x", expand=True)

        self.check_files_button = ctk.CTkButton(main_frame, text="3. Dosyaları Kontrol Et ve Alanları Eşleştir", command=self.check_files, state="disabled")
        self.check_files_button.pack(pady=(10, 5), padx=10, fill="x")

        naming_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        naming_frame.pack(pady=5, padx=10, fill="x")

        self.select_naming_button = ctk.CTkButton(naming_frame, text="4. Dosya Adı İçin Sütun Seç", width=180, command=self.select_naming_columns, state="disabled")
        self.select_naming_button.pack(side="left", padx=(0, 10))

        self.naming_label = ctk.CTkLabel(naming_frame, text="İsimlendirme sütunları seçilmedi...", text_color="gray")
        self.naming_label.pack(side="left", fill="x", expand=True)

        fields_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        fields_frame.pack(pady=5, padx=10, fill="both", expand=True)

        fields_frame.grid_columnconfigure(0, weight=1)
        fields_frame.grid_columnconfigure(1, weight=1)
        fields_frame.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(fields_frame, text="Excel Sütunları").grid(row=0, column=0, padx=5, pady=(0, 5))
        ctk.CTkLabel(fields_frame, text="Word Yer Tutucuları ({{...}})").grid(row=0, column=1, padx=5, pady=(0, 5))

        self.excel_cols_textbox = ctk.CTkTextbox(fields_frame, wrap="none", height=100)
        self.excel_cols_textbox.grid(row=1, column=0, padx=(0, 5), sticky="nsew")

        self.word_fields_textbox = ctk.CTkTextbox(fields_frame, wrap="none", height=100)
        self.word_fields_textbox.grid(row=1, column=1, padx=(5, 0), sticky="nsew")

        self.excel_cols_textbox.configure(state="disabled")
        self.word_fields_textbox.configure(state="disabled")

        self.process_button = ctk.CTkButton(main_frame, text="5. Taslakları Oluştur", height=40, command=self.process_creation, state="disabled")
        self.process_button.pack(pady=(10, 10), padx=10, fill="x")

        self.status_label = ctk.CTkLabel(main_frame, text="", font=ctk.CTkFont(size=12))
        self.status_label.pack(pady=(0, 10))

    def select_excel(self):
        file_path = filedialog.askopenfilename(
            title="Excel Veri Dosyasını Seçin",
            filetypes=[("Excel Dosyaları", "*.xlsx *.xls")]
        )
        if file_path:
            self.excel_path = file_path
            self.excel_label.configure(text=os.path.basename(file_path), text_color="white")
            self.update_check_button_state()
            self.naming_columns = []
            self.naming_label.configure(text="İsimlendirme sütunları seçilmedi...", text_color="gray")
            self.select_naming_button.configure(state="disabled")
            self.process_button.configure(state="disabled")
            self.excel_columns = []
            self.excel_cols_textbox.configure(state="normal")
            self.excel_cols_textbox.delete("1.0", "end")
            self.excel_cols_textbox.configure(state="disabled")


    def select_word(self):
        file_path = filedialog.askopenfilename(
            title="Word Şablon Dosyasını Seçin",
            filetypes=[("Word Belgeleri", "*.docx")]
        )
        if file_path:
            if not file_path.lower().endswith(".docx"):
                messagebox.showerror("Hata", "Lütfen .docx formatında bir Word dosyası seçin.")
                return
            self.word_path = file_path
            self.word_label.configure(text=os.path.basename(file_path), text_color="white")
            self.update_check_button_state()
            self.process_button.configure(state="disabled")
            self.placeholders = []
            self.word_fields_textbox.configure(state="normal")
            self.word_fields_textbox.delete("1.0", "end")
            self.word_fields_textbox.configure(state="disabled")

    def update_check_button_state(self):
        if self.excel_path and self.word_path:
            self.check_files_button.configure(state="normal")
        else:
            self.check_files_button.configure(state="disabled")

    def find_placeholders(self, document):
        placeholders = set()
        for para in document.paragraphs:
            found = re.findall(r"\{\{([^}]+)\}\}", para.text)
            placeholders.update(found) 
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found = re.findall(r"\{\{([^}]+)\}\}", para.text)
                        placeholders.update(found)
        return sorted(list(placeholders)) 

    def check_files(self):
        self.status_label.configure(text="Dosyalar kontrol ediliyor...", text_color="yellow")
        self.process_button.configure(state="disabled")
        self.select_naming_button.configure(state="disabled") 
        try:
            df = pd.read_excel(self.excel_path, nrows=0) 
            self.excel_columns = sorted([str(col) for col in df.columns])

            self.excel_cols_textbox.configure(state="normal")
            self.excel_cols_textbox.delete("1.0", "end")
            self.excel_cols_textbox.insert("1.0", "\n".join(self.excel_columns))
            self.excel_cols_textbox.configure(state="disabled")

            doc = docx.Document(self.word_path)
            self.placeholders = self.find_placeholders(doc)

            self.word_fields_textbox.configure(state="normal")
            self.word_fields_textbox.delete("1.0", "end")
            self.word_fields_textbox.insert("1.0", "\n".join(self.placeholders))
            self.word_fields_textbox.configure(state="disabled")

            if not self.placeholders:
                self.status_label.configure(text="HATA: Word şablonunda {{Sutun_Adi}} gibi bir yer tutucu bulunamadı.", text_color="red")
                messagebox.showerror("Hata", "Word şablonunda hiç {{...}} formatında yer tutucu bulunamadı.")
                return

            excel_set = set(self.excel_columns)
            word_set = set(self.placeholders)
            matches = excel_set.intersection(word_set)
            missing_in_excel = word_set - excel_set 

            if not matches:
                self.status_label.configure(text="HATA: Word yer tutucuları ile Excel sütunları arasında eşleşme yok.", text_color="red")
                messagebox.showwarning("Eşleşme Yok", "Word şablonundaki {{...}} yer tutucuları ile Excel'deki sütun başlıkları arasında hiçbir eşleşme bulunamadı.\nLütfen dosyaları kontrol edin.")
                status_text = "HATA: Eşleşme yok."
                color = "red"
            else:
                status_text = f"Kontrol başarılı. {len(matches)} adet alan eşleşti."
                color = "lightgreen"

            if missing_in_excel:
                status_text += f" UYARI: {len(missing_in_excel)} Word alanı Excel'de yok!"
                color = "orange"
                messagebox.showwarning("Eksik Sütunlar", f"Word şablonundaki şu yer tutucular Excel dosyasında bulunamadı:\n\n{', '.join(missing_in_excel)}\n\nBu alanlar değiştirilmeden bırakılacaktır.")

            if matches: 
                self.select_naming_button.configure(state="normal")
                self.status_label.configure(text=status_text + " Şimdi dosya adı için sütun seçin.", text_color=color)
            else:
                 self.status_label.configure(text=status_text, text_color=color) 

        except Exception as e:
            self.status_label.configure(text=f"Bir hata oluştu: {e}", text_color="red")
            messagebox.showerror("Hata", f"Dosyalar okunurken veya kontrol edilirken bir hata oluştu:\n{e}")
            traceback.print_exc() 
            self.select_naming_button.configure(state="disabled") 

    def select_naming_columns(self):
        if not self.excel_columns:
            messagebox.showerror("Hata", "Önce Excel dosyasını seçip kontrol etmelisin.")
            return

        dialog = ctk.CTkToplevel(self)
        dialog.title("Dosya Adı Sütunlarını Seç")
        dialog.geometry("400x450")
        dialog.transient(self)
        dialog.grab_set() 

        ctk.CTkLabel(dialog, text="Dosya adını oluşturacak sütunları sırayla seçin:", font=ctk.CTkFont(weight="bold")).pack(pady=10)
        ctk.CTkLabel(dialog, text="(Seçim sırası dosya adındaki sırayı belirler)").pack(pady=(0,5))


        scrollable_frame = ctk.CTkScrollableFrame(dialog, height=300)
        scrollable_frame.pack(fill="x", expand=False, padx=15)

        self.checkbox_vars = {}
        ordered_columns = self.naming_columns + [col for col in self.excel_columns if col not in self.naming_columns]

        for col in ordered_columns:
            var = ctk.StringVar(value="off")
            if col in self.naming_columns:
                var.set("on")
            cb = ctk.CTkCheckBox(scrollable_frame, text=col, variable=var, onvalue="on", offvalue="off")
            cb.pack(anchor="w", padx=10, pady=2)
            self.checkbox_vars[col] = var 

        def confirm_selection():
            selected_cols_in_order = []
            for col in ordered_columns: 
                 if self.checkbox_vars[col].get() == "on":
                     selected_cols_in_order.append(col)

            self.naming_columns = selected_cols_in_order 

            if self.naming_columns:
                display_text = " -> ".join(self.naming_columns) 
                if len(display_text) > 50: # 
                    display_text = display_text[:47] + "..."
                self.naming_label.configure(text=f"Sıra: {display_text}", text_color="white")
                self.process_button.configure(state="normal") 
                self.status_label.configure(text="İsimlendirme sütunları seçildi. Taslakları oluşturabilirsiniz.", text_color="lightgreen")
            else:
                self.naming_label.configure(text="İsimlendirme sütunları seçilmedi...", text_color="gray")
                self.process_button.configure(state="disabled") 
                self.status_label.configure(text="Dosya adı için sütun seçmelisiniz.", text_color="orange")
            dialog.destroy() 

        button_frame = ctk.CTkFrame(dialog, fg_color="transparent")
        button_frame.pack(pady=10, fill="x")

        confirm_button = ctk.CTkButton(button_frame, text="Onayla", command=confirm_selection)
        confirm_button.pack(side="right", padx=15)

        cancel_button = ctk.CTkButton(button_frame, text="İptal", command=dialog.destroy, fg_color="gray")
        cancel_button.pack(side="right", padx=5)

    def replace_text_in_paragraph(self, paragraph, replacements):
        for key, value in replacements.items():
            placeholder = "{{" + key + "}}"
            replacement_value = str(value) if value is not None else ""

            if placeholder in paragraph.text:
                inline = paragraph.runs
                full_text = "".join(r.text for r in inline)
                new_text = full_text.replace(placeholder, replacement_value)
                for i in range(len(inline)):
                    p = paragraph._p
                    p.remove(inline[i]._r)
                paragraph.add_run(new_text)


    def process_creation(self):
        if not self.excel_path or not self.word_path:
            messagebox.showerror("Hata", "Lütfen önce Excel ve Word dosyalarını seçin.")
            return

        if not self.excel_columns:
             messagebox.showerror("Hata", "Lütfen önce 'Dosyaları Kontrol Et' butonuna basın.")
             return

        if not self.naming_columns:
            messagebox.showerror("Hata", "Lütfen önce dosya adı için sütun seçin.")
            return

        output_dir = filedialog.askdirectory(title="Oluşturulan belgeler nereye kaydedilsin?")
        if not output_dir:
            self.status_label.configure(text="İşlem iptal edildi.", text_color="gray")
            return

        try:
            self.status_label.configure(text="Excel verisi okunuyor...", text_color="yellow")
            self.update_idletasks()

            df = pd.read_excel(self.excel_path, keep_default_na=False, dtype=str)

            self.status_label.configure(text=f"İşlem başladı... {len(df)} adet belge oluşturuluyor...", text_color="yellow")
            self.update_idletasks()

            word_base_name = os.path.splitext(os.path.basename(self.word_path))[0]
            matched_placeholders = set(self.excel_columns).intersection(set(self.placeholders))

            for index, row in df.iterrows():
                document = docx.Document(self.word_path)

                replacements = {col: row[col] for col in matched_placeholders}

                for para in document.paragraphs:
                    self.replace_text_in_paragraph(para, replacements)

                for table in document.tables:
                    for r in table.rows:
                        for cell in r.cells:
                            for para in cell.paragraphs:
                                self.replace_text_in_paragraph(para, replacements)

                try:
                    name_parts = [row[col].strip() for col in self.naming_columns]
                    name_parts = [part for part in name_parts if part]
                    if name_parts:
                        filename_base = "_".join(name_parts) 
                    else: 
                        filename_base = f"kayit_{index+1:03d}"

                    filename_base = re.sub(r'[\\/*?:"<>|\[\]]', "", filename_base)
                    max_len = 150 
                    if len(filename_base) > max_len:
                         filename_base = filename_base[:max_len]

                    if not filename_base or filename_base.strip('.') == '':
                        filename_base = f"kayit_{index+1:03d}"

                except KeyError as key_ex: 
                     print(f"HATA: Satır {index+1} için dosya adı oluşturulamadı. Sütun bulunamadı: {key_ex}")
                     messagebox.showerror("İsimlendirme Hatası", f"Dosya adı için seçilen '{key_ex}' sütunu Excel dosyasında bulunamadı.\nLütfen dosyaları tekrar kontrol edin.")
                     self.status_label.configure(text=f"İsimlendirme hatası: Sütun '{key_ex}' bulunamadı.", text_color="red")
                     return 
                except Exception as name_ex: 
                     print(f"Uyarı: Satır {index+1} için dosya adı oluşturulurken hata: {name_ex}")
                     filename_base = f"kayit_{index+1:03d}"


                output_filename = f"{word_base_name}_{filename_base}.docx"
                output_filepath = os.path.join(output_dir, output_filename)

                try:
                    document.save(output_filepath)
                except Exception as save_ex:
                    print(f"HATA: Dosya kaydedilemedi: {output_filepath}")
                    print(f"Detay: {save_ex}")

                    try:
                        simple_filename = f"{word_base_name}_kayit_{index+1:03d}.docx"
                        simple_filepath = os.path.join(output_dir, simple_filename)
                        print(f"Basit isimle deneniyor: {simple_filepath}")
                        document.save(simple_filepath)
                    except Exception as simple_save_ex:
                        print(f"Basit isimle kaydetme de BAŞARISIZ: {simple_save_ex}")
                        messagebox.showwarning("Kaydetme Hatası", f"'{output_filename}' dosyası kaydedilemedi.\nDosya adı geçersiz olabilir veya yazma izni olmayabilir.\nSatır {index+1} atlandı.")


                if (index + 1) % 20 == 0: 
                    self.status_label.configure(text=f"{index+1}/{len(df)} belge oluşturuldu...", text_color="yellow")
                    self.update_idletasks() 

            self.status_label.configure(text=f"İşlem tamamlandı! {len(df)} belge için işlem yapıldı.", text_color="lightgreen")
            messagebox.showinfo("Başarılı", f"İşlem tamamlandı.\n\n'{output_dir}' klasörüne belgeler oluşturuldu.")

        except Exception as e:
            self.status_label.configure(text=f"Bir hata oluştu: {e}", text_color="red")
            messagebox.showerror("Hata", f"Belge oluşturma sırasında bir hata oluştu:\n{e}")
            traceback.print_exc()
